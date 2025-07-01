import hashlib
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches

# Poznati maliciozni MD5 hash-evi
malicious_hashes = {
    "44d88612fea8a8f36de82e1278abb02f",
    "6f5902ac237024bdd0c176cb93063dc4",
    "098f6bcd4621d373cade4e832627b4f6" #kada ovaj MD5 hes dodamo u Temp dobijamo izlaz da je pronadjen zlonamerni fajl,
                                        #
}

# Direktorijum koji se skenira
scan_dir = "C:\\Windows\\Temp"

# Kreiranje Word dokumenta
document = Document()
document.add_heading('Izveštaj o skeniranju fajlova', 0)

# Vreme početka skeniranja
start_time = datetime.now()
document.add_paragraph(f"Datum i vreme skeniranja: {start_time.strftime('%d.%m.%Y. %H:%M:%S')}")
document.add_paragraph(f"Direktorijum za skeniranje: {scan_dir}")
document.add_paragraph("Poznati maliciozni MD5 hash-evi:")
for h in malicious_hashes:
    document.add_paragraph(h, style='List Bullet')

document.add_paragraph("\nRezultati skeniranja:")

# Funkcija za računanje MD5 hasha
def get_md5(file_path):
    try:
        with open(file_path, "rb") as f:
            file_hash = hashlib.md5()
            while chunk := f.read(4096):
                file_hash.update(chunk)
            return file_hash.hexdigest()
    except Exception as e:
        return None

# Skeniranje direktorijuma
found_malicious = []
scanned_files = 0

for root, _, files in os.walk(scan_dir):
    for file in files:
        full_path = os.path.join(root, file)
        file_hash = get_md5(full_path)
        scanned_files += 1
        if file_hash in malicious_hashes:
            found_malicious.append((full_path, file_hash))
            print(f"[ALERT] Malicious file detected: {full_path} (MD5: {file_hash})")
            document.add_paragraph(f"[Zlonamerni fajl] {full_path} (MD5: {file_hash})", style='List Number')
if not found_malicious:
    document.add_paragraph("Nema zlonamernih fajlova!")
# Rezime
document.add_paragraph("\nStatistika:")
document.add_paragraph(f"Ukupno skeniranih fajlova: {scanned_files}")
document.add_paragraph(f"Broj detektovanih zlonamernih fajlova: {len(found_malicious)}")

# Vreme završetka
end_time = datetime.now()
document.add_paragraph(f"Vreme završetka: {end_time.strftime('%d.%m.%Y. %H:%M:%S')}")
duration = end_time - start_time
document.add_paragraph(f"Trajanje skeniranja: {duration}")

# Snimanje dokumenta
report_filename = f"Izvestaj_skeniranja_{start_time.strftime('%Y%m%d_%H%M%S')}.docx"
document.save(report_filename)
print(f"\n[INFO] Izveštaj je sačuvan kao: {report_filename}")


